#!/usr/bin/env python3
"""
Multi-format knowledge ingestion pipeline for RAG.

Supported formats:
- PDF (.pdf)
- Word (.docx)
- Excel (.xlsx, .xlsm)
- CSV (.csv, .tsv)
- JSON (.json)
- Markdown (.md)
- Plain text (.txt)
- HTML (.html, .htm)

What this script does:
1. Recursively scans an input directory.
2. Parses files into normalized document sections.
3. Cleans and normalizes extracted text.
4. Chunks content using structure-aware and token-aware splitting.
5. Builds embeddings for each chunk.
6. Writes artifacts to disk:
   - normalized_documents.jsonl
   - chunks.jsonl
   - embeddings.npy
   - qdrant_upsert_payload.jsonl

Design goals:
- Single-file, easy to run, easy to modify.
- Minimal framework lock-in.
- Good metadata for later retrieval, filtering, citation, and governance.

Install example:
    pip install pypdf python-docx pandas openpyxl beautifulsoup4 markdownify sentence-transformers qdrant-client numpy

Example usage:
    python rag_multiformat_ingestion_pipeline.py \
        --input_dir ./knowledge_base \
        --output_dir ./rag_output \
        --embedding_model BAAI/bge-m3 \
        --chunk_size 700 \
        --chunk_overlap 120

Optional Qdrant upload:
    python rag_multiformat_ingestion_pipeline.py \
        --input_dir ./knowledge_base \
        --output_dir ./rag_output \
        --qdrant_url http://localhost:6333 \
        --qdrant_collection rag_chunks
"""

from __future__ import annotations

import argparse
import csv
import dataclasses
import hashlib
import html
import json
import logging
import os
import re
import sys
import traceback
from collections import Counter
from dataclasses import dataclass, asdict, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, Iterable, Iterator, List, Optional, Sequence, Tuple

import numpy as np
import pandas as pd
try:
    from bs4 import BeautifulSoup
except ImportError:
    BeautifulSoup = None

try:
    from markdownify import markdownify as html_to_markdown
except ImportError:
    html_to_markdown = None

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    from sentence_transformers import SentenceTransformer
except ImportError:
    SentenceTransformer = None

# Optional imports handled gracefully.
try:
    import docx
except ImportError:
    docx = None

try:
    from qdrant_client import QdrantClient
    from qdrant_client.http.models import Distance, FieldCondition, Filter, MatchAny, PointStruct, VectorParams
except ImportError:
    QdrantClient = None
    Distance = None
    FieldCondition = None
    Filter = None
    MatchAny = None
    PointStruct = None
    VectorParams = None


LOGGER = logging.getLogger("rag_ingestion")
UTC = timezone.utc
SUPPORTED_SUFFIXES = {
    ".pdf",
    ".docx",
    ".xlsx",
    ".xlsm",
    ".csv",
    ".tsv",
    ".json",
    ".md",
    ".txt",
    ".html",
    ".htm",
}


# =========================
# Data models
# =========================

@dataclass
class RawSection:
    heading: str
    content: str
    page: Optional[int] = None
    order: int = 0
    section_path: List[str] = field(default_factory=list)


@dataclass
class NormalizedDocument:
    doc_id: str
    source_path: str
    source_name: str
    doc_type: str
    title: str
    language: str
    created_at: Optional[str]
    updated_at: Optional[str]
    checksum: str
    metadata: Dict[str, Any]
    sections: List[RawSection]
    summary: str


@dataclass
class SyncManifest:
    generated_at: str
    input_dir: str
    output_dir: str
    collection_name: str
    embedding_model: str
    documents: List[Dict[str, Any]]
    chunk_count: int


@dataclass
class ChunkRecord:
    chunk_id: str
    doc_id: str
    source_path: str
    source_name: str
    doc_type: str
    title: str
    section: str
    section_path: List[str]
    page_start: Optional[int]
    page_end: Optional[int]
    chunk_index: int
    content: str
    content_for_embedding: str
    summary: str
    tags: List[str]
    keywords: List[str]
    language: str
    created_at: Optional[str]
    updated_at: Optional[str]
    metadata: Dict[str, Any]


# =========================
# Utilities
# =========================

def utc_now_iso() -> str:
    return datetime.now(tz=UTC).isoformat()


def sha256_text(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8", errors="ignore")).hexdigest()



def sha256_file(path: Path) -> str:
    hasher = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            hasher.update(chunk)
    return hasher.hexdigest()



def stable_id(*parts: str, length: int = 16) -> str:
    joined = "||".join(parts)
    return hashlib.sha256(joined.encode("utf-8", errors="ignore")).hexdigest()[:length]



def normalize_whitespace(text: str) -> str:
    text = text.replace("\u00a0", " ")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[\t\f\v]+", " ", text)
    text = re.sub(r" +", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()



def remove_repetitive_lines(text: str, min_repeat: int = 3) -> str:
    lines = [line.strip() for line in text.splitlines()]
    counter = Counter([line for line in lines if line])
    filtered = [line for line in lines if not (line and counter[line] >= min_repeat and len(line) < 120)]
    return normalize_whitespace("\n".join(filtered))



def clean_text(text: str) -> str:
    text = html.unescape(text)
    text = normalize_whitespace(text)
    text = remove_repetitive_lines(text)
    # Common OCR / extraction cleanup.
    text = re.sub(r"(?<=\w)-\n(?=\w)", "", text)
    text = re.sub(r"\n(?=[,.;:!?])", "", text)
    return normalize_whitespace(text)



def guess_language(text: str) -> str:
    # Lightweight heuristic. Replace with langdetect/fasttext if needed.
    if re.search(r"[\u4e00-\u9fff]", text):
        return "zh"
    return "en"



def extract_keywords(text: str, top_k: int = 12) -> List[str]:
    tokens = re.findall(r"[A-Za-z][A-Za-z0-9_\-]{2,}|[\u4e00-\u9fff]{2,}", text.lower())
    stopwords = {
        "the", "and", "for", "with", "that", "this", "from", "you", "your", "are", "was", "were",
        "have", "has", "had", "not", "can", "will", "shall", "into", "their", "there", "about", "what",
        "which", "when", "where", "how", "why", "then", "than", "also", "use", "using", "used",
        "以及", "我们", "你们", "进行", "相关", "一个", "可以", "需要", "用于", "其中", "说明", "通过",
    }
    filtered = [t for t in tokens if t not in stopwords and len(t) >= 2]
    counts = Counter(filtered)
    return [word for word, _ in counts.most_common(top_k)]



def build_summary(text: str, max_chars: int = 280) -> str:
    text = clean_text(text)
    if len(text) <= max_chars:
        return text
    split = re.split(r"(?<=[。！？.!?])\s+|\n\n", text)
    summary = ""
    for part in split:
        if not part:
            continue
        candidate = (summary + " " + part).strip()
        if len(candidate) > max_chars:
            break
        summary = candidate
    if summary:
        return summary
    return text[: max_chars - 1].rstrip() + "…"



def write_jsonl(path: Path, rows: Iterable[Dict[str, Any]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8") as f:
        for row in rows:
            f.write(json.dumps(row, ensure_ascii=False) + "\n")



def load_manifest(path: Path) -> Optional[Dict[str, Any]]:
    if not path.exists():
        return None
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        LOGGER.warning("Failed to load manifest from %s", path)
        return None



def build_manifest(
    input_dir: Path,
    output_dir: Path,
    documents: Sequence[NormalizedDocument],
    chunks: Sequence[ChunkRecord],
    embedding_model: str,
    collection_name: str,
) -> SyncManifest:
    rows = []
    for doc in documents:
        rows.append(
            {
                "doc_id": doc.doc_id,
                "source_path": doc.source_path,
                "source_name": doc.source_name,
                "checksum": doc.checksum,
                "doc_type": doc.doc_type,
                "updated_at": doc.updated_at,
            }
        )
    return SyncManifest(
        generated_at=utc_now_iso(),
        input_dir=str(input_dir),
        output_dir=str(output_dir),
        collection_name=collection_name,
        embedding_model=embedding_model,
        documents=rows,
        chunk_count=len(chunks),
    )



def persist_manifest(output_dir: Path, manifest: SyncManifest) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)
    manifest_path = output_dir / "ingestion_manifest.json"
    manifest_path.write_text(json.dumps(asdict(manifest), ensure_ascii=False, indent=2), encoding="utf-8")



def collect_stale_doc_ids(previous_manifest: Optional[Dict[str, Any]], documents: Sequence[NormalizedDocument]) -> List[str]:
    if not previous_manifest:
        return []
    previous_by_source = {
        row.get("source_path"): row
        for row in previous_manifest.get("documents", [])
        if isinstance(row, dict) and row.get("source_path") and row.get("doc_id")
    }
    current_by_source = {doc.source_path: doc for doc in documents}
    stale_doc_ids: List[str] = []
    for source_path, previous in previous_by_source.items():
        current = current_by_source.get(source_path)
        if current is None or current.checksum != previous.get("checksum"):
            stale_doc_ids.append(str(previous["doc_id"]))
    return sorted(set(stale_doc_ids))


def read_text_file(path: Path) -> str:
    encodings = ["utf-8", "utf-8-sig", "gb18030", "latin-1"]
    last_error: Optional[Exception] = None
    for enc in encodings:
        try:
            return path.read_text(encoding=enc)
        except Exception as exc:
            last_error = exc
    raise RuntimeError(f"Failed to read text file {path}: {last_error}")



def safe_slug(value: str) -> str:
    value = value.strip().lower()
    value = re.sub(r"[^a-z0-9\u4e00-\u9fff]+", "-", value)
    value = re.sub(r"-+", "-", value).strip("-")
    return value or "doc"


# =========================
# Parsers
# =========================

class BaseParser:
    def parse(self, path: Path) -> NormalizedDocument:
        raise NotImplementedError


class PDFParser(BaseParser):
    def parse(self, path: Path) -> NormalizedDocument:
        if PdfReader is None:
            raise RuntimeError('pypdf is not installed.')
        reader = PdfReader(str(path))
        sections: List[RawSection] = []
        all_text_parts: List[str] = []

        for idx, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            text = clean_text(text)
            if not text:
                continue
            heading = f"Page {idx + 1}"
            sections.append(
                RawSection(
                    heading=heading,
                    content=text,
                    page=idx + 1,
                    order=idx,
                    section_path=[heading],
                )
            )
            all_text_parts.append(text)

        raw_text = "\n\n".join(all_text_parts)
        title = path.stem
        language = guess_language(raw_text)
        checksum = sha256_file(path)
        summary = build_summary(raw_text)

        return NormalizedDocument(
            doc_id=f"pdf_{stable_id(str(path), checksum)}",
            source_path=str(path),
            source_name=path.name,
            doc_type="pdf",
            title=title,
            language=language,
            created_at=None,
            updated_at=datetime.fromtimestamp(path.stat().st_mtime, tz=UTC).isoformat(),
            checksum=checksum,
            metadata={"page_count": len(reader.pages)},
            sections=sections,
            summary=summary,
        )


class DOCXParser(BaseParser):
    HEADING_STYLE_PREFIX = "Heading"

    def parse(self, path: Path) -> NormalizedDocument:
        if docx is None:
            raise RuntimeError("python-docx is not installed.")

        document = docx.Document(str(path))
        sections: List[RawSection] = []
        current_heading_stack: List[str] = [path.stem]
        buffer: List[str] = []
        order = 0

        def flush_buffer() -> None:
            nonlocal order, buffer
            content = clean_text("\n".join(buffer))
            if content:
                heading = current_heading_stack[-1] if current_heading_stack else path.stem
                sections.append(
                    RawSection(
                        heading=heading,
                        content=content,
                        page=None,
                        order=order,
                        section_path=current_heading_stack.copy(),
                    )
                )
                order += 1
            buffer = []

        for para in document.paragraphs:
            text = clean_text(para.text)
            if not text:
                continue
            style_name = getattr(para.style, "name", "") or ""
            if style_name.startswith(self.HEADING_STYLE_PREFIX):
                flush_buffer()
                level_match = re.search(r"(\d+)$", style_name)
                level = int(level_match.group(1)) if level_match else 1
                current_heading_stack = current_heading_stack[:level - 1] + [text]
            else:
                buffer.append(text)

        flush_buffer()
        raw_text = "\n\n".join(section.content for section in sections)
        title = path.stem
        language = guess_language(raw_text)
        checksum = sha256_file(path)
        summary = build_summary(raw_text)

        return NormalizedDocument(
            doc_id=f"docx_{stable_id(str(path), checksum)}",
            source_path=str(path),
            source_name=path.name,
            doc_type="docx",
            title=title,
            language=language,
            created_at=None,
            updated_at=datetime.fromtimestamp(path.stat().st_mtime, tz=UTC).isoformat(),
            checksum=checksum,
            metadata={},
            sections=sections,
            summary=summary,
        )


class ExcelParser(BaseParser):
    def parse(self, path: Path) -> NormalizedDocument:
        excel = pd.ExcelFile(path)
        sections: List[RawSection] = []
        all_text: List[str] = []

        for sheet_order, sheet_name in enumerate(excel.sheet_names):
            df = excel.parse(sheet_name=sheet_name)
            text = self._sheet_to_text(df, sheet_name)
            text = clean_text(text)
            if not text:
                continue
            sections.append(
                RawSection(
                    heading=sheet_name,
                    content=text,
                    page=None,
                    order=sheet_order,
                    section_path=[sheet_name],
                )
            )
            all_text.append(text)

        raw_text = "\n\n".join(all_text)
        title = path.stem
        language = guess_language(raw_text)
        checksum = sha256_file(path)
        summary = build_summary(raw_text)

        return NormalizedDocument(
            doc_id=f"excel_{stable_id(str(path), checksum)}",
            source_path=str(path),
            source_name=path.name,
            doc_type="excel",
            title=title,
            language=language,
            created_at=None,
            updated_at=datetime.fromtimestamp(path.stat().st_mtime, tz=UTC).isoformat(),
            checksum=checksum,
            metadata={"sheet_names": excel.sheet_names},
            sections=sections,
            summary=summary,
        )

    @staticmethod
    def _sheet_to_text(df: pd.DataFrame, sheet_name: str) -> str:
        if df.empty:
            return ""
        df = df.dropna(axis=0, how="all").dropna(axis=1, how="all")
        if df.empty:
            return ""
        columns = [str(c).strip() for c in df.columns]
        lines = [f"Sheet: {sheet_name}"]
        for row_idx, row in df.iterrows():
            kvs = []
            for col, value in zip(columns, row.tolist()):
                if pd.isna(value):
                    continue
                value_text = str(value).strip()
                if not value_text:
                    continue
                kvs.append(f"{col}: {value_text}")
            if kvs:
                lines.append(f"Row {row_idx + 1}: " + "; ".join(kvs))
        return "\n".join(lines)


class CSVParser(BaseParser):
    def parse(self, path: Path) -> NormalizedDocument:
        sep = "\t" if path.suffix.lower() == ".tsv" else ","
        df = pd.read_csv(path, sep=sep, dtype=str, keep_default_na=False)
        text = ExcelParser._sheet_to_text(df, path.stem)
        text = clean_text(text)
        checksum = sha256_file(path)
        return NormalizedDocument(
            doc_id=f"csv_{stable_id(str(path), checksum)}",
            source_path=str(path),
            source_name=path.name,
            doc_type="csv",
            title=path.stem,
            language=guess_language(text),
            created_at=None,
            updated_at=datetime.fromtimestamp(path.stat().st_mtime, tz=UTC).isoformat(),
            checksum=checksum,
            metadata={},
            sections=[RawSection(heading=path.stem, content=text, order=0, section_path=[path.stem])],
            summary=build_summary(text),
        )


class JSONParser(BaseParser):
    def parse(self, path: Path) -> NormalizedDocument:
        raw = json.loads(read_text_file(path))
        text = self._json_to_text(raw)
        text = clean_text(text)
        checksum = sha256_file(path)
        return NormalizedDocument(
            doc_id=f"json_{stable_id(str(path), checksum)}",
            source_path=str(path),
            source_name=path.name,
            doc_type="json",
            title=path.stem,
            language=guess_language(text),
            created_at=None,
            updated_at=datetime.fromtimestamp(path.stat().st_mtime, tz=UTC).isoformat(),
            checksum=checksum,
            metadata={},
            sections=[RawSection(heading=path.stem, content=text, order=0, section_path=[path.stem])],
            summary=build_summary(text),
        )

    def _json_to_text(self, obj: Any, prefix: str = "") -> str:
        lines: List[str] = []
        if isinstance(obj, dict):
            for key, value in obj.items():
                new_prefix = f"{prefix}.{key}" if prefix else str(key)
                lines.append(self._json_to_text(value, new_prefix))
        elif isinstance(obj, list):
            for idx, item in enumerate(obj):
                new_prefix = f"{prefix}[{idx}]"
                lines.append(self._json_to_text(item, new_prefix))
        else:
            value = "" if obj is None else str(obj)
            lines.append(f"{prefix}: {value}")
        return "\n".join([line for line in lines if line])


class MarkdownParser(BaseParser):
    def parse(self, path: Path) -> NormalizedDocument:
        text = read_text_file(path)
        sections = split_markdown_sections(text, fallback_title=path.stem)
        raw_text = "\n\n".join(section.content for section in sections)
        checksum = sha256_file(path)
        return NormalizedDocument(
            doc_id=f"md_{stable_id(str(path), checksum)}",
            source_path=str(path),
            source_name=path.name,
            doc_type="markdown",
            title=path.stem,
            language=guess_language(raw_text),
            created_at=None,
            updated_at=datetime.fromtimestamp(path.stat().st_mtime, tz=UTC).isoformat(),
            checksum=checksum,
            metadata={},
            sections=sections,
            summary=build_summary(raw_text),
        )


class TextParser(BaseParser):
    def parse(self, path: Path) -> NormalizedDocument:
        text = clean_text(read_text_file(path))
        checksum = sha256_file(path)
        sections = [RawSection(heading=path.stem, content=text, order=0, section_path=[path.stem])]
        return NormalizedDocument(
            doc_id=f"txt_{stable_id(str(path), checksum)}",
            source_path=str(path),
            source_name=path.name,
            doc_type="text",
            title=path.stem,
            language=guess_language(text),
            created_at=None,
            updated_at=datetime.fromtimestamp(path.stat().st_mtime, tz=UTC).isoformat(),
            checksum=checksum,
            metadata={},
            sections=sections,
            summary=build_summary(text),
        )


class HTMLParser(BaseParser):
    def parse(self, path: Path) -> NormalizedDocument:
        if BeautifulSoup is None or html_to_markdown is None:
            raise RuntimeError('beautifulsoup4 and markdownify are required for HTML ingestion.')
        raw_html = read_text_file(path)
        soup = BeautifulSoup(raw_html, "html.parser")

        title = soup.title.get_text(strip=True) if soup.title else path.stem

        for tag in soup(["script", "style", "noscript"]):
            tag.extract()

        body_html = str(soup.body or soup)
        markdown = html_to_markdown(body_html)
        sections = split_markdown_sections(markdown, fallback_title=title)
        raw_text = "\n\n".join(section.content for section in sections)
        checksum = sha256_file(path)

        return NormalizedDocument(
            doc_id=f"html_{stable_id(str(path), checksum)}",
            source_path=str(path),
            source_name=path.name,
            doc_type="html",
            title=title,
            language=guess_language(raw_text),
            created_at=None,
            updated_at=datetime.fromtimestamp(path.stat().st_mtime, tz=UTC).isoformat(),
            checksum=checksum,
            metadata={},
            sections=sections,
            summary=build_summary(raw_text),
        )


# =========================
# Markdown / structure handling
# =========================

def split_markdown_sections(text: str, fallback_title: str) -> List[RawSection]:
    text = clean_text(text)
    if not text:
        return [RawSection(heading=fallback_title, content="", order=0, section_path=[fallback_title])]

    lines = text.splitlines()
    sections: List[RawSection] = []
    heading_stack: List[str] = [fallback_title]
    buffer: List[str] = []
    order = 0

    def flush() -> None:
        nonlocal order, buffer
        content = clean_text("\n".join(buffer))
        if content:
            sections.append(
                RawSection(
                    heading=heading_stack[-1],
                    content=content,
                    order=order,
                    section_path=heading_stack.copy(),
                )
            )
            order += 1
        buffer = []

    for line in lines:
        m = re.match(r"^(#{1,6})\s+(.*)$", line.strip())
        if m:
            flush()
            level = len(m.group(1))
            heading = clean_text(m.group(2))
            heading_stack = heading_stack[:level - 1] + [heading]
        else:
            buffer.append(line)

    flush()
    if not sections:
        sections = [RawSection(heading=fallback_title, content=text, order=0, section_path=[fallback_title])]
    return sections


# =========================
# Chunking
# =========================

def estimate_tokens(text: str) -> int:
    # Approximation sufficient for ingestion chunking.
    return max(1, int(len(text) / 4))


class TextChunker:
    def __init__(self, chunk_size: int = 700, chunk_overlap: int = 120):
        if chunk_overlap >= chunk_size:
            raise ValueError("chunk_overlap must be smaller than chunk_size")
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def split_document(self, doc: NormalizedDocument) -> List[ChunkRecord]:
        chunks: List[ChunkRecord] = []
        chunk_index = 0

        for section in doc.sections:
            if not section.content.strip():
                continue

            section_chunks = self._split_text(section.content)
            for local_idx, piece in enumerate(section_chunks):
                section_name = section.heading or doc.title
                keywords = extract_keywords(piece)
                tags = list({doc.doc_type, safe_slug(section_name), *keywords[:5]})
                content_for_embedding = self._build_embedding_text(doc, section, piece)
                chunk_id = f"{doc.doc_id}_chunk_{chunk_index:05d}"

                chunks.append(
                    ChunkRecord(
                        chunk_id=chunk_id,
                        doc_id=doc.doc_id,
                        source_path=doc.source_path,
                        source_name=doc.source_name,
                        doc_type=doc.doc_type,
                        title=doc.title,
                        section=section_name,
                        section_path=section.section_path,
                        page_start=section.page,
                        page_end=section.page,
                        chunk_index=chunk_index,
                        content=piece,
                        content_for_embedding=content_for_embedding,
                        summary=doc.summary,
                        tags=tags,
                        keywords=keywords,
                        language=doc.language,
                        created_at=doc.created_at,
                        updated_at=doc.updated_at,
                        metadata={
                            **doc.metadata,
                            "local_section_chunk_index": local_idx,
                        },
                    )
                )
                chunk_index += 1

        return chunks

    def _build_embedding_text(self, doc: NormalizedDocument, section: RawSection, content: str) -> str:
        parts = [
            f"Title: {doc.title}",
            f"Document type: {doc.doc_type}",
            f"Section path: {' > '.join(section.section_path)}",
            f"Section heading: {section.heading}",
            f"Content: {content}",
        ]
        return "\n".join(parts)

    def _split_text(self, text: str) -> List[str]:
        paragraphs = [p.strip() for p in re.split(r"\n\n+", text) if p.strip()]
        if not paragraphs:
            return []

        chunks: List[str] = []
        current: List[str] = []
        current_tokens = 0

        for paragraph in paragraphs:
            p_tokens = estimate_tokens(paragraph)
            if p_tokens > self.chunk_size:
                if current:
                    chunk_text = "\n\n".join(current).strip()
                    if chunk_text:
                        chunks.append(chunk_text)
                    current = []
                    current_tokens = 0
                sentence_chunks = self._split_large_paragraph(paragraph)
                chunks.extend(sentence_chunks)
                continue

            if current_tokens + p_tokens <= self.chunk_size:
                current.append(paragraph)
                current_tokens += p_tokens
                continue

            chunk_text = "\n\n".join(current).strip()
            if chunk_text:
                chunks.append(chunk_text)
            current = self._overlap_tail(current)
            current_tokens = estimate_tokens("\n\n".join(current)) if current else 0
            current.append(paragraph)
            current_tokens += p_tokens

        if current:
            chunk_text = "\n\n".join(current).strip()
            if chunk_text:
                chunks.append(chunk_text)
        return [clean_text(c) for c in chunks if clean_text(c)]

    def _split_large_paragraph(self, paragraph: str) -> List[str]:
        sentences = re.split(r"(?<=[\u3002\uff01\uff1f.!?])\s+", paragraph)
        parts: List[str] = []
        current: List[str] = []
        current_tokens = 0
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
            s_tokens = estimate_tokens(sentence)
            if s_tokens > self.chunk_size:
                if current:
                    parts.append(" ".join(current).strip())
                    current = []
                    current_tokens = 0
                parts.extend(self._split_oversized_text(sentence))
                continue
            if current_tokens + s_tokens <= self.chunk_size:
                current.append(sentence)
                current_tokens += s_tokens
            else:
                if current:
                    parts.append(" ".join(current).strip())
                current = [sentence]
                current_tokens = s_tokens
        if current:
            parts.append(" ".join(current).strip())
        return [clean_text(part) for part in self._with_text_overlap(parts) if clean_text(part)]

    def _split_oversized_text(self, text: str) -> List[str]:
        words = re.findall(r"\S+", text)
        if not words:
            return []
        parts: List[str] = []
        current: List[str] = []
        current_tokens = 0
        for word in words:
            word_tokens = estimate_tokens(word)
            if current and current_tokens + word_tokens > self.chunk_size:
                parts.append(" ".join(current).strip())
                current = self._word_overlap_tail(current)
                current_tokens = estimate_tokens(" ".join(current)) if current else 0
            if word_tokens > self.chunk_size:
                fragments = self._split_oversized_word(word)
                for fragment in fragments:
                    fragment_tokens = estimate_tokens(fragment)
                    if current and current_tokens + fragment_tokens > self.chunk_size:
                        parts.append(" ".join(current).strip())
                        current = self._word_overlap_tail(current)
                        current_tokens = estimate_tokens(" ".join(current)) if current else 0
                    current.append(fragment)
                    current_tokens += fragment_tokens
                continue
            current.append(word)
            current_tokens += word_tokens
        if current:
            parts.append(" ".join(current).strip())
        return [clean_text(part) for part in self._with_text_overlap(parts) if clean_text(part)]

    def _split_oversized_word(self, word: str) -> List[str]:
        max_chars = max(8, self.chunk_size * 4)
        return [word[idx: idx + max_chars] for idx in range(0, len(word), max_chars)]

    def _word_overlap_tail(self, words: List[str]) -> List[str]:
        tail: List[str] = []
        tokens = 0
        for word in reversed(words):
            word_tokens = estimate_tokens(word)
            if tokens + word_tokens > self.chunk_overlap:
                break
            tail.insert(0, word)
            tokens += word_tokens
        return tail

    def _overlap_tail(self, paragraphs: List[str]) -> List[str]:
        if not paragraphs:
            return []
        tail: List[str] = []
        tokens = 0
        for paragraph in reversed(paragraphs):
            p_tokens = estimate_tokens(paragraph)
            if tokens + p_tokens > self.chunk_overlap:
                break
            tail.insert(0, paragraph)
            tokens += p_tokens
        return tail

    def _with_text_overlap(self, parts: List[str]) -> List[str]:
        if not parts:
            return []
        overlapped: List[str] = []
        previous = ""
        overlap_chars = max(40, self.chunk_overlap * 4)
        for part in parts:
            part = part.strip()
            if not part:
                continue
            prefix = previous[-overlap_chars:].strip() if previous else ""
            combined = part if not prefix or part.startswith(prefix) else f"{prefix}\n\n{part}"
            overlapped.append(combined)
            previous = part
        return overlapped



# =========================
# Embedding and storage
# =========================

class EmbeddingIndexer:
    def __init__(self, model_name: str, batch_size: int = 32, normalize_embeddings: bool = True, fallback_dim: int = 256):
        self.requested_model_name = model_name
        self.batch_size = batch_size
        self.normalize_embeddings = normalize_embeddings
        self.fallback_dim = fallback_dim
        self.model = None
        self.model_name = model_name
        if SentenceTransformer is None:
            LOGGER.warning("sentence-transformers is not installed. Falling back to hashing embeddings.")
            self.model_name = f"hashing-fallback:{fallback_dim}"
            return
        allow_remote = os.environ.get("RAG_ALLOW_REMOTE_MODEL_DOWNLOAD", "").strip().lower() in {"1", "true", "yes"}
        try:
            if allow_remote:
                self.model = SentenceTransformer(model_name)
            else:
                self.model = SentenceTransformer(model_name, local_files_only=True)
        except TypeError:
            if allow_remote:
                self.model = SentenceTransformer(model_name)
            else:
                LOGGER.warning("Embedding model %s is not available in local cache. Falling back to hashing embeddings.", model_name)
                self.model_name = f"hashing-fallback:{fallback_dim}"
                self.model = None
        except Exception as exc:
            LOGGER.warning("Failed to load embedding model %s, falling back to hashing embeddings: %s", model_name, exc)
            self.model_name = f"hashing-fallback:{fallback_dim}"
            self.model = None

    def encode(self, texts: Sequence[str]) -> np.ndarray:
        if self.model is None:
            return self._encode_hashing(texts)
        vectors = self.model.encode(
            list(texts),
            batch_size=self.batch_size,
            show_progress_bar=True,
            normalize_embeddings=self.normalize_embeddings,
            convert_to_numpy=True,
        )
        if not isinstance(vectors, np.ndarray):
            vectors = np.asarray(vectors)
        return vectors.astype(np.float32)

    def _encode_hashing(self, texts: Sequence[str]) -> np.ndarray:
        rows = [self._hash_text(text) for text in texts]
        return np.asarray(rows, dtype=np.float32)

    def _hash_text(self, text: str) -> np.ndarray:
        vector = np.zeros(self.fallback_dim, dtype=np.float32)
        terms = re.findall(r"[A-Za-z0-9_\-]{2,}|[\u4e00-\u9fff]{1,}", (text or '').lower())
        if not terms:
            return vector
        for term in terms:
            digest = hashlib.sha256(term.encode('utf-8')).digest()
            index = int.from_bytes(digest[:4], 'little') % self.fallback_dim
            sign = 1.0 if digest[4] % 2 == 0 else -1.0
            vector[index] += sign
        if self.normalize_embeddings:
            norm = float(np.linalg.norm(vector))
            if norm > 0:
                vector = vector / norm
        return vector


class QdrantUploader:
    def __init__(self, url: str, collection_name: str):
        if QdrantClient is None:
            raise RuntimeError("qdrant-client is not installed.")
        self.url = url
        self.collection_name = collection_name
        self.client = QdrantClient(url=url)

    def ensure_collection(self, vector_size: int) -> None:
        existing = [c.name for c in self.client.get_collections().collections]
        if self.collection_name not in existing:
            self.client.create_collection(
                collection_name=self.collection_name,
                vectors_config=VectorParams(size=vector_size, distance=Distance.COSINE),
            )

    def delete_doc_ids(self, doc_ids: Sequence[str]) -> None:
        if not doc_ids:
            return
        if Filter is None or FieldCondition is None or MatchAny is None:
            raise RuntimeError("qdrant-client does not provide delete filter models.")
        self.client.delete(
            collection_name=self.collection_name,
            points_selector=Filter(
                must=[
                    FieldCondition(key="doc_id", match=MatchAny(any=list(doc_ids))),
                ]
            ),
        )

    def upload(self, chunks: Sequence[ChunkRecord], vectors: np.ndarray, batch_size: int = 128) -> None:
        if len(chunks) != len(vectors):
            raise ValueError("chunks and vectors must have the same length")
        self.ensure_collection(vector_size=int(vectors.shape[1]))
        points: List[PointStruct] = []
        for chunk, vector in zip(chunks, vectors):
            payload = asdict(chunk)
            points.append(
                PointStruct(
                    id=chunk.chunk_id,
                    vector=vector.tolist(),
                    payload=payload,
                )
            )
        for start in range(0, len(points), batch_size):
            self.client.upsert(
                collection_name=self.collection_name,
                points=points[start:start + batch_size],
            )


# =========================
# Pipeline
# =========================

def get_parser_for_path(path: Path) -> BaseParser:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return PDFParser()
    if suffix == ".docx":
        return DOCXParser()
    if suffix in {".xlsx", ".xlsm"}:
        return ExcelParser()
    if suffix in {".csv", ".tsv"}:
        return CSVParser()
    if suffix == ".json":
        return JSONParser()
    if suffix == ".md":
        return MarkdownParser()
    if suffix in {".txt"}:
        return TextParser()
    if suffix in {".html", ".htm"}:
        return HTMLParser()
    raise ValueError(f"Unsupported file type: {suffix}")



def collect_files(input_dir: Path) -> List[Path]:
    files = [p for p in input_dir.rglob("*") if p.is_file() and p.suffix.lower() in SUPPORTED_SUFFIXES]
    files.sort()
    return files



def normalize_document(path: Path) -> NormalizedDocument:
    parser = get_parser_for_path(path)
    return parser.parse(path)



def persist_outputs(
    output_dir: Path,
    documents: Sequence[NormalizedDocument],
    chunks: Sequence[ChunkRecord],
    embeddings: np.ndarray,
) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)

    normalized_rows = []
    for doc in documents:
        row = asdict(doc)
        row["sections"] = [asdict(section) for section in doc.sections]
        normalized_rows.append(row)
    write_jsonl(output_dir / "normalized_documents.jsonl", normalized_rows)

    chunk_rows = [asdict(chunk) for chunk in chunks]
    write_jsonl(output_dir / "chunks.jsonl", chunk_rows)

    np.save(output_dir / "embeddings.npy", embeddings)

    qdrant_rows = []
    for chunk, vector in zip(chunks, embeddings):
        qdrant_rows.append(
            {
                "id": chunk.chunk_id,
                "vector": vector.tolist(),
                "payload": asdict(chunk),
            }
        )
    write_jsonl(output_dir / "qdrant_upsert_payload.jsonl", qdrant_rows)



def run_pipeline(args: argparse.Namespace) -> int:
    input_dir = Path(args.input_dir).resolve()
    output_dir = Path(args.output_dir).resolve()
    manifest_path = output_dir / "ingestion_manifest.json"
    previous_manifest = load_manifest(manifest_path)

    if not input_dir.exists() or not input_dir.is_dir():
        raise FileNotFoundError(f"Input directory does not exist or is not a directory: {input_dir}")

    files = collect_files(input_dir)
    if not files:
        LOGGER.warning("No supported files found in %s", input_dir)
        return 0

    LOGGER.info("Found %d supported files", len(files))

    documents: List[NormalizedDocument] = []
    failed_files: List[Dict[str, str]] = []

    for path in files:
        LOGGER.info("Parsing: %s", path)
        try:
            doc = normalize_document(path)
            documents.append(doc)
        except Exception as exc:
            LOGGER.exception("Failed to parse %s", path)
            failed_files.append({"path": str(path), "error": str(exc)})

    if not documents:
        LOGGER.error("All files failed to parse.")
        write_jsonl(output_dir / "failed_files.jsonl", failed_files)
        return 1

    chunker = TextChunker(chunk_size=args.chunk_size, chunk_overlap=args.chunk_overlap)
    chunks: List[ChunkRecord] = []
    for doc in documents:
        chunks.extend(chunker.split_document(doc))

    if not chunks:
        LOGGER.error("No chunks were produced.")
        write_jsonl(output_dir / "failed_files.jsonl", failed_files)
        return 1

    LOGGER.info("Generated %d chunks", len(chunks))

    embedder = EmbeddingIndexer(
        model_name=args.embedding_model,
        batch_size=args.embedding_batch_size,
        normalize_embeddings=not args.disable_normalize_embeddings,
    )
    embeddings = embedder.encode([chunk.content_for_embedding for chunk in chunks])
    LOGGER.info("Embeddings shape: %s", embeddings.shape)

    persist_outputs(output_dir, documents, chunks, embeddings)

    if failed_files:
        write_jsonl(output_dir / "failed_files.jsonl", failed_files)

    if args.qdrant_url and args.qdrant_collection:
        LOGGER.info("Uploading to Qdrant: %s / %s", args.qdrant_url, args.qdrant_collection)
        uploader = QdrantUploader(url=args.qdrant_url, collection_name=args.qdrant_collection)
        if args.sync_qdrant:
            stale_doc_ids = collect_stale_doc_ids(previous_manifest, documents)
            if stale_doc_ids:
                LOGGER.info("Deleting %d stale document(s) from Qdrant before upsert", len(stale_doc_ids))
                uploader.delete_doc_ids(stale_doc_ids)
        uploader.upload(chunks, embeddings, batch_size=args.qdrant_batch_size)

    manifest = build_manifest(
        input_dir=input_dir,
        output_dir=output_dir,
        documents=documents,
        chunks=chunks,
        embedding_model=embedder.model_name,
        collection_name=args.qdrant_collection or "",
    )
    persist_manifest(output_dir, manifest)

    LOGGER.info("Done. Output written to %s", output_dir)
    return 0


# =========================
# CLI
# =========================

def parse_args(argv: Optional[Sequence[str]] = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Multi-format knowledge ingestion pipeline for RAG")
    parser.add_argument("--input_dir", required=True, help="Directory containing source files")
    parser.add_argument("--output_dir", required=True, help="Directory for normalized outputs and embeddings")
    parser.add_argument(
        "--embedding_model",
        default="BAAI/bge-m3",
        help="SentenceTransformer embedding model name",
    )
    parser.add_argument("--embedding_batch_size", type=int, default=32, help="Embedding batch size")
    parser.add_argument("--chunk_size", type=int, default=700, help="Approximate chunk size in tokens")
    parser.add_argument("--chunk_overlap", type=int, default=120, help="Approximate overlap in tokens")
    parser.add_argument(
        "--disable_normalize_embeddings",
        action="store_true",
        help="Disable L2 normalization for embeddings",
    )
    parser.add_argument("--qdrant_url", default="", help="Optional Qdrant URL, e.g. http://localhost:6333")
    parser.add_argument("--qdrant_collection", default="", help="Optional Qdrant collection name")
    parser.add_argument("--qdrant_batch_size", type=int, default=128, help="Qdrant upsert batch size")
    parser.add_argument("--sync_qdrant", action="store_true", help="Delete stale document versions from Qdrant before upsert")
    parser.add_argument("--log_level", default="INFO", help="Logging level")
    return parser.parse_args(argv)



def configure_logging(level: str) -> None:
    logging.basicConfig(
        level=getattr(logging, level.upper(), logging.INFO),
        format="%(asctime)s | %(levelname)s | %(name)s | %(message)s",
    )



def main(argv: Optional[Sequence[str]] = None) -> int:
    args = parse_args(argv)
    configure_logging(args.log_level)
    try:
        return run_pipeline(args)
    except Exception as exc:
        LOGGER.error(r"Pipeline failed: %s", exc)
        LOGGER.debug(r"Traceback:\n%s", traceback.format_exc())
        return 1


if __name__ == "__main__":
    sys.exit(main())
